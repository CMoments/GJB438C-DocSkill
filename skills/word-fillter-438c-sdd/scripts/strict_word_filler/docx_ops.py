from __future__ import annotations

from copy import deepcopy
import re

from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from .errors import StrictTemplateError
from .models import BuildPlan, DynamicSectionRule, TableFillRule

CAPTION_PLACEHOLDER_RE = re.compile(r"^(?P<kind>[表图])(?P<chapter>\d+|X)-X(?P<rest>.*)$")


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from iter_container_paragraphs(cell)


def iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        yield from _iter_table_paragraphs(table)


def iter_all_paragraphs(doc: DocumentObject):
    yield from iter_container_paragraphs(doc)
    for section in doc.sections:
        yield from iter_container_paragraphs(section.header)
        yield from iter_container_paragraphs(section.footer)


def iter_body_non_toc_paragraphs(doc: DocumentObject):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "toc" in style_name:
            continue
        yield paragraph


def _pick_style_run(paragraph: Paragraph):
    for run in paragraph.runs:
        if run.text:
            return run
    return paragraph.runs[0] if paragraph.runs else None


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        style_run = _pick_style_run(paragraph)
        if style_run is None:
            paragraph.add_run(text)
            return
        style_run.text = text
        style_run_element = style_run._r
        for run in paragraph.runs:
            if run._r is style_run_element:
                continue
            run.text = ""
        return
    paragraph.add_run(text)


def _paragraph_lines(text: str) -> list[str]:
    raw_lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    lines = [line for line in raw_lines if line.strip()]
    if not lines:
        raise StrictTemplateError("写入文本不能为空")
    return lines


def set_paragraphs_from_text(paragraph: Paragraph, text: str) -> None:
    lines = _paragraph_lines(text)
    set_paragraph_text(paragraph, lines[0])
    cursor = paragraph
    for line in lines[1:]:
        new_paragraph = _clone_after(cursor, paragraph)
        set_paragraph_text(new_paragraph, line)
        cursor = new_paragraph


def apply_exact_replacement(doc: DocumentObject, old_text: str, new_text: str, expected_matches: int, key: str) -> None:
    matches = [paragraph for paragraph in iter_all_paragraphs(doc) if paragraph.text == old_text]
    if len(matches) != expected_matches:
        raise StrictTemplateError(
            f"{key} 期望命中 {expected_matches} 个精确段落，实际 {len(matches)} 个：{old_text}"
        )
    for paragraph in matches:
        set_paragraph_text(paragraph, new_text)


def apply_paragraph_replacement(doc: DocumentObject, anchor_text: str, new_text: str, expected_matches: int, key: str) -> None:
    matches = [paragraph for paragraph in iter_all_paragraphs(doc) if anchor_text in paragraph.text]
    if len(matches) != expected_matches:
        raise StrictTemplateError(
            f"{key} 期望命中 {expected_matches} 个段落，实际 {len(matches)} 个：{anchor_text}"
        )
    for paragraph in matches:
        set_paragraphs_from_text(paragraph, new_text)


def find_single_body_paragraph(doc: DocumentObject, text: str, *, exact: bool = False) -> Paragraph:
    matches = []
    for paragraph in iter_body_non_toc_paragraphs(doc):
        if exact and paragraph.text == text:
            matches.append(paragraph)
        elif not exact and text in paragraph.text:
            matches.append(paragraph)
    if len(matches) != 1:
        mode = "精确匹配" if exact else "包含匹配"
        raise StrictTemplateError(f"模板中 {mode} `{text}` 的正文段落应为 1 个，实际为 {len(matches)} 个")
    return matches[0]


def paragraph_index(doc: DocumentObject, paragraph: Paragraph) -> int:
    for index, candidate in enumerate(doc.paragraphs):
        if candidate._p is paragraph._p:
            return index
    raise StrictTemplateError("无法在正文段落列表中定位目标段落")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def _clone_after(anchor: Paragraph, prototype: Paragraph) -> Paragraph:
    new_p = deepcopy(prototype._p)
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def _heading_separator(prototype_heading_text: str) -> str:
    stripped = prototype_heading_text.lstrip("0123456789.X")
    if stripped.startswith(" "):
        return " "
    return ""


def _content_blocks(body: str) -> list[str]:
    return _paragraph_lines(body)


def _is_heading1(paragraph: Paragraph) -> bool:
    if paragraph.style is None:
        return False
    style_name = paragraph.style.name.lower()
    return "heading 1" in style_name or "标题 1" in style_name


def _row_texts(row: _Row) -> tuple[str, ...]:
    return tuple(cell.text.strip() for cell in row.cells)


def _find_table_by_locator(doc: DocumentObject, locator_rows: tuple[tuple[str, ...], ...], key: str) -> Table:
    matches: list[Table] = []
    for table in doc.tables:
        if len(table.rows) < len(locator_rows):
            continue
        if all(_row_texts(table.rows[index]) == locator_row for index, locator_row in enumerate(locator_rows)):
            matches.append(table)
    if len(matches) != 1:
        raise StrictTemplateError(f"{key} 期望命中 1 张表，实际 {len(matches)} 张")
    return matches[0]


def _insert_row_before(table: Table, row_index: int, prototype_row_index: int) -> _Row:
    prototype = table.rows[prototype_row_index]
    new_tr = deepcopy(prototype._tr)
    if row_index >= len(table.rows):
        table._tbl.append(new_tr)
    else:
        table.rows[row_index]._tr.addprevious(new_tr)
    return _Row(new_tr, table)


def _prepare_cell(cell: _Cell) -> Paragraph:
    if not cell.paragraphs:
        return cell.add_paragraph()
    first = cell.paragraphs[0]
    for paragraph in list(cell.paragraphs[1:]):
        delete_paragraph(paragraph)
    clear_paragraph(first)
    return first


def _set_cell_text(cell: _Cell, text: str) -> None:
    paragraph = _prepare_cell(cell)
    set_paragraphs_from_text(paragraph, text)


def apply_table_fill(doc: DocumentObject, rule: TableFillRule) -> None:
    table = _find_table_by_locator(doc, rule.locator_rows, rule.key)
    expected_columns = len(rule.columns)
    if len(table.columns) != expected_columns:
        raise StrictTemplateError(
            f"{rule.key} 期望 {expected_columns} 列，模板实际 {len(table.columns)} 列"
        )
    if len(table.rows) <= rule.data_start_row:
        raise StrictTemplateError(f"{rule.key} 的 data_start_row 超出模板行数")
    if rule.preserve_tail_rows > len(table.rows) - rule.data_start_row:
        raise StrictTemplateError(f"{rule.key} 的 preserve_tail_rows 非法")

    current_data_rows = len(table.rows) - rule.data_start_row - rule.preserve_tail_rows
    if current_data_rows <= 0:
        raise StrictTemplateError(f"{rule.key} 没有可用的数据行模板")

    while len(rule.rows) > len(table.rows) - rule.data_start_row - rule.preserve_tail_rows:
        insert_at = len(table.rows) - rule.preserve_tail_rows
        prototype_row_index = insert_at - 1
        if prototype_row_index < rule.data_start_row:
            raise StrictTemplateError(f"{rule.key} 无法确定可克隆的数据行原型")
        _insert_row_before(table, insert_at, prototype_row_index)

    for row_offset, row_values in enumerate(rule.rows):
        target_row = table.rows[rule.data_start_row + row_offset]
        if len(target_row.cells) != expected_columns:
            raise StrictTemplateError(f"{rule.key} 的目标行列数异常")
        for cell_index, value in enumerate(row_values):
            _set_cell_text(target_row.cells[cell_index], value)


def apply_caption_numbering(doc: DocumentObject) -> None:
    current_chapter = 0
    counters: dict[tuple[str, int], int] = {}

    for paragraph in doc.paragraphs:
        if _is_heading1(paragraph):
            current_chapter += 1
            continue

        text = paragraph.text.strip()
        match = CAPTION_PLACEHOLDER_RE.match(text)
        if match is None:
            continue

        explicit_chapter = match.group("chapter")
        if explicit_chapter == "X":
            if current_chapter == 0:
                raise StrictTemplateError(f"图表标题 `{text}` 出现在任何一级章节之前")
            chapter = current_chapter
        else:
            chapter = int(explicit_chapter)

        key = (match.group("kind"), chapter)
        counters[key] = counters.get(key, 0) + 1
        set_paragraph_text(paragraph, f"{match.group('kind')}{chapter}-{counters[key]}{match.group('rest')}")


def apply_dynamic_section(doc: DocumentObject, rule: DynamicSectionRule) -> None:
    parent_paragraph = find_single_body_paragraph(doc, rule.parent_title, exact=True)
    summary_paragraph = find_single_body_paragraph(doc, rule.summary_anchor, exact=False)
    end_before_paragraph = find_single_body_paragraph(doc, rule.end_before_title, exact=True)
    prototype_heading = find_single_body_paragraph(doc, rule.prototype_title_anchor, exact=True)
    prototype_body = find_single_body_paragraph(doc, rule.prototype_body_anchor, exact=False)

    parent_index = paragraph_index(doc, parent_paragraph)
    summary_index = paragraph_index(doc, summary_paragraph)
    end_before_index = paragraph_index(doc, end_before_paragraph)
    prototype_heading_index = paragraph_index(doc, prototype_heading)
    prototype_body_index = paragraph_index(doc, prototype_body)

    if not (parent_index < summary_index < end_before_index):
        raise StrictTemplateError(f"{rule.parent_id} 的父标题、摘要和结束锚点顺序不合法")

    cleanup_paragraphs: list[Paragraph] = []
    for anchor in rule.cleanup_anchors:
        cleanup_paragraph = find_single_body_paragraph(doc, anchor, exact=False)
        cleanup_index = paragraph_index(doc, cleanup_paragraph)
        if not (summary_index < cleanup_index < end_before_index):
            raise StrictTemplateError(f"{rule.parent_id} 的清理锚点不在目标章节范围内: {anchor}")
        cleanup_paragraphs.append(cleanup_paragraph)

    prototype_title_inside = summary_index < prototype_heading_index < end_before_index
    prototype_body_inside = summary_index < prototype_body_index < end_before_index

    paragraphs_to_remove = cleanup_paragraphs[:]
    if prototype_title_inside:
        paragraphs_to_remove.append(prototype_heading)
    if prototype_body_inside:
        paragraphs_to_remove.append(prototype_body)

    if rule.cleanup_to_end:
        if not paragraphs_to_remove:
            raise StrictTemplateError(f"{rule.parent_id} 要求从清理起点删到章节结束，但未找到清理锚点")
        cleanup_start = min(paragraph_index(doc, paragraph) for paragraph in paragraphs_to_remove)
        paragraphs_to_remove = list(doc.paragraphs[cleanup_start:end_before_index])

    if paragraphs_to_remove:
        insertion_before_index = min(paragraph_index(doc, paragraph) for paragraph in paragraphs_to_remove)
    else:
        insertion_before_index = end_before_index

    if insertion_before_index <= summary_index:
        raise StrictTemplateError(f"{rule.parent_id} 的插入点非法")

    insertion_after = doc.paragraphs[insertion_before_index - 1]
    separator = _heading_separator(prototype_heading.text)
    cursor = insertion_after
    for item in rule.items:
        new_heading = _clone_after(cursor, prototype_heading)
        set_paragraph_text(new_heading, f"{item.section_id}{separator}{item.title}")
        cursor = new_heading
        for block in _content_blocks(item.body):
            new_body = _clone_after(cursor, prototype_body)
            set_paragraph_text(new_body, block)
            cursor = new_body

    unique_to_remove = []
    seen = set()
    for paragraph in paragraphs_to_remove:
        marker = id(paragraph._p)
        if marker in seen:
            continue
        seen.add(marker)
        unique_to_remove.append(paragraph)

    for paragraph in sorted(unique_to_remove, key=lambda item: paragraph_index(doc, item), reverse=True):
        delete_paragraph(paragraph)


def mark_toc_for_update(doc: DocumentObject) -> None:
    settings_element = doc.settings.element
    update_field = settings_element.find(qn("w:updateFields"))
    if update_field is None:
        update_field = OxmlElement("w:updateFields")
        settings_element.append(update_field)
    update_field.set(qn("w:val"), "true")

    for field in doc._element.body.iter(qn("w:fldChar")):
        if field.get(qn("w:fldCharType")) == "begin":
            field.set(qn("w:dirty"), "true")


def apply_plan(doc: DocumentObject, plan: BuildPlan) -> None:
    for rule in plan.exact_replacements:
        apply_exact_replacement(doc, rule.old_text, rule.new_text, rule.expected_matches, rule.key)

    for rule in plan.dynamic_sections:
        apply_dynamic_section(doc, rule)

    for rule in plan.paragraph_replacements:
        apply_paragraph_replacement(doc, rule.anchor_text, rule.new_text, rule.expected_matches, rule.key)

    for rule in plan.tables:
        apply_table_fill(doc, rule)

    apply_caption_numbering(doc)
    mark_toc_for_update(doc)
